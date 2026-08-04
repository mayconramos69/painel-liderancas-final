import os, csv, io
from datetime import datetime
from functools import wraps

import psycopg
from psycopg import errors
from psycopg.rows import dict_row
from flask import Flask, render_template, request, redirect, url_for, session, flash, Response, send_file, abort
from werkzeug.security import generate_password_hash, check_password_hash
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4, landscape
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Inches, Pt

app = Flask(__name__, static_folder="static", template_folder="templates")
app.secret_key = os.environ.get("SECRET_KEY", "chave-temporaria")
app.config["MAX_CONTENT_LENGTH"] = 6 * 1024 * 1024

DATABASE_URL = os.environ.get("DATABASE_URL")
if DATABASE_URL and DATABASE_URL.startswith("postgres://"):
    DATABASE_URL = DATABASE_URL.replace("postgres://", "postgresql://", 1)

ADMIN_USER = os.environ.get("ADMIN_USER", "mayconramos2026")
ADMIN_PASS = os.environ.get("ADMIN_PASS", "26511076mj")



@app.template_filter("whatsapp_numero")
def whatsapp_numero(telefone):
    """Transforma telefone brasileiro em número compatível com wa.me."""
    digitos = "".join(ch for ch in str(telefone or "") if ch.isdigit())
    if not digitos:
        return ""
    digitos = digitos.lstrip("0")
    if digitos.startswith("55") and len(digitos) >= 12:
        return digitos
    if len(digitos) in (10, 11):
        return "55" + digitos
    return digitos


UF_TITULO = {
    "01": "São Paulo", "02": "Minas Gerais", "03": "Rio de Janeiro",
    "04": "Rio Grande do Sul", "05": "Bahia", "06": "Paraná",
    "07": "Ceará", "08": "Pernambuco", "09": "Santa Catarina",
    "10": "Goiás", "11": "Maranhão", "12": "Paraíba",
    "13": "Pará", "14": "Espírito Santo", "15": "Piauí",
    "16": "Rio Grande do Norte", "17": "Alagoas", "18": "Mato Grosso",
    "19": "Mato Grosso do Sul", "20": "Distrito Federal", "21": "Sergipe",
    "22": "Amazonas", "23": "Rondônia", "24": "Acre",
    "25": "Amapá", "26": "Roraima", "27": "Tocantins",
    "28": "Exterior"
}

def normalizar_titulo(valor):
    return "".join(ch for ch in str(valor or "") if ch.isdigit())


def analisar_titulo_eleitor(valor):
    """Valida a estrutura e os dígitos verificadores do título eleitoral.

    Isso confirma que o número tem formato matematicamente válido. A situação
    cadastral do eleitor só pode ser confirmada nos serviços oficiais do TSE.
    """
    digitos = normalizar_titulo(valor)
    if not digitos:
        return {"digitos": "", "valido": None, "uf": "Não informado", "codigo_uf": ""}
    if len(digitos) != 12 or len(set(digitos)) == 1:
        return {"digitos": digitos, "valido": False, "uf": "Origem não identificada", "codigo_uf": ""}

    codigo_uf = digitos[8:10]
    uf = UF_TITULO.get(codigo_uf, "Origem não identificada")
    numeros = [int(x) for x in digitos]

    # 1º DV: oito primeiros algarismos, pesos 2 a 9.
    resto1 = sum(numeros[i] * (i + 2) for i in range(8)) % 11
    dv1 = 0 if resto1 == 10 else resto1

    # 2º DV: código da UF e primeiro DV, pesos 7, 8 e 9.
    resto2 = (numeros[8] * 7 + numeros[9] * 8 + dv1 * 9) % 11
    dv2 = 0 if resto2 == 10 else resto2

    valido = numeros[10] == dv1 and numeros[11] == dv2 and codigo_uf in UF_TITULO
    return {"digitos": digitos, "valido": valido, "uf": uf, "codigo_uf": codigo_uf}

@app.template_filter("titulo_info")
def titulo_info(valor):
    return analisar_titulo_eleitor(valor)

def normalizar(texto):
    return " ".join((texto or "").strip().lower().split())


def db():
    if not DATABASE_URL:
        raise RuntimeError("DATABASE_URL não configurada no Render.")
    return psycopg.connect(DATABASE_URL, row_factory=dict_row)

def init_db():
    conn = db()
    cur = conn.cursor()

    cur.execute("""
    CREATE TABLE IF NOT EXISTS usuarios (
        id SERIAL PRIMARY KEY,
        nome TEXT NOT NULL,
        usuario TEXT UNIQUE NOT NULL,
        senha_hash TEXT NOT NULL,
        telefone TEXT,
        email TEXT,
        municipio TEXT,
        bairro TEXT,
        zona_regiao TEXT,
        perfil TEXT NOT NULL DEFAULT 'lideranca',
        status TEXT NOT NULL DEFAULT 'pendente',
        pode_trabalho INTEGER NOT NULL DEFAULT 0,
        created_at TEXT NOT NULL
    )
    """)


    # Atualização segura do banco existente. Apenas adiciona o que estiver faltando.
    for comando in [
        "ALTER TABLE usuarios ADD COLUMN IF NOT EXISTS telefone TEXT",
        "ALTER TABLE usuarios ADD COLUMN IF NOT EXISTS email TEXT",
        "ALTER TABLE usuarios ADD COLUMN IF NOT EXISTS municipio TEXT",
        "ALTER TABLE usuarios ADD COLUMN IF NOT EXISTS bairro TEXT",
        "ALTER TABLE usuarios ADD COLUMN IF NOT EXISTS zona_regiao TEXT",
        "ALTER TABLE usuarios ADD COLUMN IF NOT EXISTS perfil TEXT NOT NULL DEFAULT 'lideranca'",
        "ALTER TABLE usuarios ADD COLUMN IF NOT EXISTS status TEXT NOT NULL DEFAULT 'pendente'",
        "ALTER TABLE usuarios ADD COLUMN IF NOT EXISTS pode_trabalho INTEGER NOT NULL DEFAULT 0",
        "ALTER TABLE usuarios ADD COLUMN IF NOT EXISTS created_at TEXT"
    ]:
        cur.execute(comando)

    cur.execute("""
    CREATE TABLE IF NOT EXISTS espontaneos (
        id SERIAL PRIMARY KEY,
        lideranca_id INTEGER NOT NULL REFERENCES usuarios(id) ON DELETE CASCADE,
        nome_completo TEXT NOT NULL,
        municipio TEXT NOT NULL,
        telefone TEXT,
        endereco_completo TEXT,
        nome_normalizado TEXT,
        telefone_normalizado TEXT,
        created_at TEXT NOT NULL,
        edit_liberado INTEGER NOT NULL DEFAULT 0
    )
    """)

    for comando in [
        "ALTER TABLE espontaneos ADD COLUMN IF NOT EXISTS telefone TEXT",
        "ALTER TABLE espontaneos ADD COLUMN IF NOT EXISTS endereco_completo TEXT",
        "ALTER TABLE espontaneos ADD COLUMN IF NOT EXISTS nome_normalizado TEXT",
        "ALTER TABLE espontaneos ADD COLUMN IF NOT EXISTS telefone_normalizado TEXT",
        "ALTER TABLE espontaneos ADD COLUMN IF NOT EXISTS created_at TEXT",
        "ALTER TABLE espontaneos ADD COLUMN IF NOT EXISTS edit_liberado INTEGER NOT NULL DEFAULT 0"
    ]:
        cur.execute(comando)

    cur.execute("""
    CREATE TABLE IF NOT EXISTS trabalho (
        id SERIAL PRIMARY KEY,
        lideranca_id INTEGER NOT NULL REFERENCES usuarios(id) ON DELETE CASCADE,
        nome TEXT NOT NULL,
        municipio TEXT NOT NULL,
        colegio TEXT,
        endereco TEXT,
        telefone TEXT,
        zona TEXT,
        secao TEXT,
        numero_titulo TEXT,
        nome_normalizado TEXT,
        telefone_normalizado TEXT,
        titulo_normalizado TEXT,
        titulo_valido INTEGER,
        titulo_uf TEXT,
        tipo_trabalho TEXT NOT NULL DEFAULT 'Boca de Urna',
        votos_previstos INTEGER NOT NULL DEFAULT 0,
        foto BYTEA,
        foto_mime TEXT,
        edit_liberado INTEGER NOT NULL DEFAULT 0,
        created_at TEXT NOT NULL
    )
    """)

    cur.execute("ALTER TABLE trabalho ADD COLUMN IF NOT EXISTS titulo_valido INTEGER")
    cur.execute("ALTER TABLE trabalho ADD COLUMN IF NOT EXISTS titulo_uf TEXT")
    cur.execute("ALTER TABLE trabalho ADD COLUMN IF NOT EXISTS tipo_trabalho TEXT NOT NULL DEFAULT 'Boca de Urna'")
    cur.execute("ALTER TABLE trabalho ADD COLUMN IF NOT EXISTS votos_previstos INTEGER NOT NULL DEFAULT 0")
    cur.execute("ALTER TABLE trabalho ADD COLUMN IF NOT EXISTS foto BYTEA")
    cur.execute("ALTER TABLE trabalho ADD COLUMN IF NOT EXISTS foto_mime TEXT")
    cur.execute("ALTER TABLE trabalho ADD COLUMN IF NOT EXISTS edit_liberado INTEGER NOT NULL DEFAULT 0")
    cur.execute("ALTER TABLE usuarios ADD COLUMN IF NOT EXISTS votos_estimados INTEGER NOT NULL DEFAULT 0")
    cur.execute("SELECT id FROM usuarios WHERE usuario=%s", (ADMIN_USER,))
    admin = cur.fetchone()

    if not admin:
        cur.execute("""
            INSERT INTO usuarios
            (nome, usuario, senha_hash, telefone, email, municipio, bairro, zona_regiao, perfil, status, pode_trabalho, created_at)
            VALUES (%s, %s, %s, '', '', '', '', '', 'admin', 'ativo', 1, %s)
        """, ("Administrador", ADMIN_USER, generate_password_hash(ADMIN_PASS), datetime.now().isoformat()))
    else:
        cur.execute("UPDATE usuarios SET pode_trabalho=1, status='ativo', perfil='admin' WHERE usuario=%s", (ADMIN_USER,))

    # Recalcula registros antigos para que a tela de Válidos/Inválidos fique correta.
    cur.execute("SELECT id, numero_titulo FROM trabalho")
    for registro in cur.fetchall():
        info = analisar_titulo_eleitor(registro.get("numero_titulo"))
        cur.execute(
            "UPDATE trabalho SET titulo_normalizado=%s, titulo_valido=%s, titulo_uf=%s WHERE id=%s",
            (info["digitos"], None if info["valido"] is None else int(info["valido"]), info["uf"], registro["id"])
        )

    conn.commit()
    cur.close()
    conn.close()


def login_required(f):
    @wraps(f)
    def wrapper(*args, **kwargs):
        if "user_id" not in session:
            return redirect(url_for("index"))
        return f(*args, **kwargs)
    return wrapper


def admin_required(f):
    @wraps(f)
    def wrapper(*args, **kwargs):
        if session.get("perfil") != "admin":
            return redirect(url_for("index"))
        return f(*args, **kwargs)
    return wrapper


def lideranca_required(f):
    @wraps(f)
    def wrapper(*args, **kwargs):
        if session.get("perfil") != "lideranca":
            return redirect(url_for("index"))
        return f(*args, **kwargs)
    return wrapper


@app.route("/")
def index():
    if session.get("perfil") == "admin":
        return redirect(url_for("admin"))
    if session.get("perfil") == "lideranca":
        return redirect(url_for("lideranca"))
    return render_template("index.html")


@app.route("/cadastro", methods=["POST"])
def cadastro():
    nome = request.form.get("nome", "").strip()
    usuario = request.form.get("usuario", "").strip().lower()
    senha = request.form.get("senha", "").strip()
    telefone = request.form.get("telefone", "").strip()
    email = request.form.get("email", "").strip()
    municipio = request.form.get("municipio", "").strip()
    bairro = request.form.get("bairro", "").strip()
    zona_regiao = request.form.get("zona_regiao", "").strip()

    if not nome or not usuario or not senha:
        flash("Preencha nome, usuário e senha.")
        return redirect(url_for("index"))

    conn = db()
    cur = conn.cursor()
    try:
        cur.execute("""
            INSERT INTO usuarios
            (nome, usuario, senha_hash, telefone, email, municipio, bairro, zona_regiao, perfil, status, pode_trabalho, created_at)
            VALUES (%s, %s, %s, %s, %s, %s, %s, %s, 'lideranca', 'pendente', 0, %s)
        """, (nome, usuario, generate_password_hash(senha), telefone, email, municipio, bairro, zona_regiao, datetime.now().isoformat()))
        conn.commit()
        flash("Cadastro enviado. Aguarde aprovação do administrador.")
    except errors.UniqueViolation:
        conn.rollback()
        flash("Esse usuário já existe.")
    finally:
        cur.close()
        conn.close()

    return redirect(url_for("index"))


@app.route("/login", methods=["POST"])
def login():
    usuario = request.form.get("usuario", "").strip().lower()
    senha = request.form.get("senha", "").strip()

    conn = db()
    cur = conn.cursor()
    cur.execute("SELECT * FROM usuarios WHERE usuario=%s", (usuario,))
    user = cur.fetchone()
    cur.close()
    conn.close()

    if not user or not check_password_hash(user["senha_hash"], senha):
        flash("Usuário ou senha inválidos.")
        return redirect(url_for("index"))

    if user["status"] != "ativo":
        flash("Seu acesso ainda não foi aprovado ou está bloqueado.")
        return redirect(url_for("index"))

    session["user_id"] = user["id"]
    session["nome"] = user["nome"]
    session["perfil"] = user["perfil"]

    return redirect(url_for("admin" if user["perfil"] == "admin" else "lideranca"))


@app.route("/sair")
def sair():
    session.clear()
    return redirect(url_for("index"))


@app.route("/lideranca")
@login_required
@lideranca_required
def lideranca():
    lid = session["user_id"]
    conn = db()
    cur = conn.cursor()

    cur.execute("SELECT * FROM usuarios WHERE id=%s", (lid,))
    user = cur.fetchone()

    cur.execute("SELECT * FROM espontaneos WHERE lideranca_id=%s ORDER BY id DESC", (lid,))
    esp = cur.fetchall()

    trab = []
    if user and int(user["pode_trabalho"] or 0) == 1:
        cur.execute("SELECT * FROM trabalho WHERE lideranca_id=%s ORDER BY id DESC", (lid,))
        trab = cur.fetchall()

    cur.close()
    conn.close()

    return render_template("lideranca.html", user=user, espontaneos=esp, trabalhos=trab)


@app.route("/espontaneo/novo", methods=["POST"])
@login_required
@lideranca_required
def novo_espontaneo():
    nome = request.form.get("nome_completo", "").strip()
    municipio = request.form.get("municipio", "").strip()
    telefone = request.form.get("telefone", "").strip()
    endereco = request.form.get("endereco_completo", "").strip()

    nn = normalizar(nome)
    tn = normalizar(telefone)

    conn = db()
    cur = conn.cursor()

    cur.execute("""
        SELECT u.nome AS lideranca_nome
        FROM espontaneos e
        JOIN usuarios u ON u.id=e.lideranca_id
        WHERE
            (e.telefone_normalizado != '' AND e.telefone_normalizado=%s)
            OR
            (e.nome_normalizado=%s AND LOWER(TRIM(e.municipio))=LOWER(TRIM(%s)))
        LIMIT 1
    """, (tn, nn, municipio))

    dup = cur.fetchone()

    if dup:
        cur.close()
        conn.close()
        flash(f"Cadastro bloqueado: já consta vinculado à liderança {dup['lideranca_nome']}.")
        return redirect(url_for("lideranca"))

    cur.execute("""
        INSERT INTO espontaneos
        (lideranca_id, nome_completo, municipio, telefone, endereco_completo, nome_normalizado, telefone_normalizado, created_at)
        VALUES (%s, %s, %s, %s, %s, %s, %s, %s)
    """, (session["user_id"], nome, municipio, telefone, endereco, nn, tn, datetime.now().isoformat()))

    conn.commit()
    cur.close()
    conn.close()

    flash("Cadastro espontâneo salvo.")
    return redirect(url_for("lideranca"))


@app.route("/trabalho/novo", methods=["POST"])
@login_required
@lideranca_required
def novo_trabalho():
    conn = db()
    cur = conn.cursor()

    cur.execute("SELECT pode_trabalho FROM usuarios WHERE id=%s", (session["user_id"],))
    user = cur.fetchone()

    if not user or int(user["pode_trabalho"] or 0) != 1:
        cur.close()
        conn.close()
        flash("Você ainda não tem autorização para cadastrar em Trabalho.")
        return redirect(url_for("lideranca"))

    nome = request.form.get("nome", "").strip()
    municipio = request.form.get("municipio", "").strip()
    colegio = request.form.get("colegio", "").strip()
    endereco = request.form.get("endereco", "").strip()
    telefone = request.form.get("telefone", "").strip()
    zona = request.form.get("zona", "").strip()
    secao = request.form.get("secao", "").strip()
    titulo = request.form.get("numero_titulo", "").strip()
    tipo_trabalho = request.form.get("tipo_trabalho", "Boca de Urna").strip()
    if tipo_trabalho not in ("Líder de rua", "Boca de Urna", "Trabalho de BU"):
        tipo_trabalho = "Boca de Urna"
    try:
        votos_previstos = max(0, int(request.form.get("votos_previstos", "0") or 0))
    except ValueError:
        votos_previstos = 0

    nn = normalizar(nome)
    tn = normalizar(telefone)
    titn = normalizar_titulo(titulo)
    titulo_info = analisar_titulo_eleitor(titulo)

    foto_arquivo = request.files.get("foto")
    foto = None
    foto_mime = None
    if foto_arquivo and foto_arquivo.filename:
        foto = foto_arquivo.read()
        foto_mime = foto_arquivo.mimetype or "image/jpeg"
        if len(foto) > 5 * 1024 * 1024:
            cur.close(); conn.close()
            flash("A foto deve ter no máximo 5 MB.")
            return redirect(url_for("lideranca"))

    cur.execute("""
        SELECT u.nome AS lideranca_nome
        FROM trabalho t
        JOIN usuarios u ON u.id=t.lideranca_id
        WHERE
            (t.titulo_normalizado != '' AND t.titulo_normalizado=%s)
            OR
            (t.telefone_normalizado != '' AND t.telefone_normalizado=%s)
            OR
            (t.nome_normalizado=%s AND LOWER(TRIM(t.municipio))=LOWER(TRIM(%s)))
        LIMIT 1
    """, (titn, tn, nn, municipio))

    dup = cur.fetchone()

    if dup:
        cur.close()
        conn.close()
        flash(f"Cadastro bloqueado: já consta vinculado à liderança {dup['lideranca_nome']}.")
        return redirect(url_for("lideranca"))

    cur.execute("""
        INSERT INTO trabalho
        (lideranca_id, nome, municipio, colegio, endereco, telefone, zona, secao, numero_titulo,
         nome_normalizado, telefone_normalizado, titulo_normalizado, titulo_valido, titulo_uf, tipo_trabalho, votos_previstos, foto, foto_mime, created_at)
        VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
    """, (session["user_id"], nome, municipio, colegio, endereco, telefone, zona, secao, titulo, nn, tn, titn,
          None if titulo_info["valido"] is None else int(titulo_info["valido"]), titulo_info["uf"], tipo_trabalho, votos_previstos, foto, foto_mime, datetime.now().isoformat()))

    conn.commit()
    cur.close()
    conn.close()

    flash("Cadastro de trabalho salvo.")
    return redirect(url_for("lideranca"))


@app.route("/trabalho/foto/<int:item_id>")
@login_required
def foto_trabalho(item_id):
    conn = db()
    cur = conn.cursor()
    if session.get("perfil") == "admin":
        cur.execute("SELECT foto, foto_mime FROM trabalho WHERE id=%s", (item_id,))
    else:
        cur.execute("SELECT foto, foto_mime FROM trabalho WHERE id=%s AND lideranca_id=%s", (item_id, session.get("user_id")))
    item = cur.fetchone()
    cur.close(); conn.close()
    if not item or not item["foto"]:
        abort(404)
    return Response(bytes(item["foto"]), mimetype=item["foto_mime"] or "image/jpeg", headers={"Cache-Control":"public, max-age=86400"})


@app.route("/admin")
@login_required
@admin_required
def admin():
    busca = request.args.get("busca", "").strip()
    municipio = request.args.get("municipio", "").strip()
    lideranca_id = request.args.get("lideranca_id", "").strip()
    titulo_status = request.args.get("titulo_status", "todos").strip()
    if titulo_status not in {"todos", "validos", "invalidos", "nao_informados"}:
        titulo_status = "todos"

    conn = db()
    cur = conn.cursor()

    cur.execute("""
        SELECT u.*, u.votos_estimados,
            (SELECT COUNT(*) FROM espontaneos e WHERE e.lideranca_id=u.id) AS total_espontaneos,
            (SELECT COUNT(*) FROM trabalho t WHERE t.lideranca_id=u.id) AS total_trabalho,
            ((SELECT COUNT(*) FROM espontaneos e WHERE e.lideranca_id=u.id) + (SELECT COUNT(*) FROM trabalho t WHERE t.lideranca_id=u.id)) AS total_geral
        FROM usuarios u
        WHERE u.perfil='lideranca'
        ORDER BY u.status DESC, total_geral DESC, u.nome ASC
    """)
    liderancas = cur.fetchall()

    cur.execute("""
        SELECT u.id, u.nome, u.municipio, u.status, u.pode_trabalho,
            COUNT(DISTINCT e.id) AS total_espontaneos,
            COUNT(DISTINCT t.id) AS total_trabalho,
            (COUNT(DISTINCT e.id) + COUNT(DISTINCT t.id)) AS total_geral
        FROM usuarios u
        LEFT JOIN espontaneos e ON e.lideranca_id=u.id
        LEFT JOIN trabalho t ON t.lideranca_id=u.id
        WHERE u.perfil='lideranca'
        GROUP BY u.id
        ORDER BY total_geral DESC, total_trabalho DESC, total_espontaneos DESC, u.nome ASC
        LIMIT 20
    """)
    ranking = cur.fetchall()

    cur.execute("SELECT COUNT(*) AS c FROM usuarios WHERE perfil='lideranca'")
    total_liderancas = cur.fetchone()["c"]

    cur.execute("SELECT COUNT(*) AS c FROM usuarios WHERE perfil='lideranca' AND status='pendente'")
    total_pendentes = cur.fetchone()["c"]

    cur.execute("SELECT COUNT(*) AS c FROM espontaneos")
    total_esp = cur.fetchone()["c"]

    cur.execute("SELECT COUNT(*) AS c FROM trabalho")
    total_trab = cur.fetchone()["c"]

    cur.execute("SELECT COUNT(*) AS c FROM trabalho WHERE titulo_valido=1")
    total_titulos_validos = cur.fetchone()["c"]
    cur.execute("SELECT COUNT(*) AS c FROM trabalho WHERE titulo_valido=0")
    total_titulos_invalidos = cur.fetchone()["c"]
    cur.execute("SELECT COUNT(*) AS c FROM trabalho WHERE titulo_valido IS NULL")
    total_titulos_nao_informados = cur.fetchone()["c"]

    cur.execute("SELECT COALESCE(SUM(votos_previstos), 0) AS total FROM trabalho")
    total_votos_previstos = cur.fetchone()["total"]

    cur.execute("SELECT COUNT(*) AS c FROM trabalho WHERE tipo_trabalho='Líder de rua'")
    total_lideres_rua = cur.fetchone()["c"]

    cur.execute("SELECT COUNT(*) AS c FROM trabalho WHERE tipo_trabalho IN ('Boca de Urna','Trabalho de BU') OR tipo_trabalho IS NULL OR tipo_trabalho=''")
    total_trabalho_bu = cur.fetchone()["c"]

    where_e, params_e = [], []
    where_t, params_t = [], []

    if busca:
        where_e.append("(e.nome_completo ILIKE %s OR e.telefone ILIKE %s OR u.nome ILIKE %s)")
        params_e += [f"%{busca}%", f"%{busca}%", f"%{busca}%"]

        where_t.append("(t.nome ILIKE %s OR t.telefone ILIKE %s OR t.colegio ILIKE %s OR u.nome ILIKE %s)")
        params_t += [f"%{busca}%", f"%{busca}%", f"%{busca}%", f"%{busca}%"]

    if municipio:
        where_e.append("e.municipio ILIKE %s")
        params_e.append(f"%{municipio}%")

        where_t.append("t.municipio ILIKE %s")
        params_t.append(f"%{municipio}%")

    if lideranca_id:
        where_e.append("e.lideranca_id=%s")
        params_e.append(lideranca_id)

        where_t.append("t.lideranca_id=%s")
        params_t.append(lideranca_id)

    sql_e = """
        SELECT e.*, u.nome AS lideranca_nome
        FROM espontaneos e
        JOIN usuarios u ON u.id=e.lideranca_id
    """
    if where_e:
        sql_e += " WHERE " + " AND ".join(where_e)
    sql_e += " ORDER BY e.id DESC"

    cur.execute(sql_e, params_e)
    esp = cur.fetchall()

    sql_t = """
        SELECT t.*, u.nome AS lideranca_nome
        FROM trabalho t
        JOIN usuarios u ON u.id=t.lideranca_id
    """
    if where_t:
        sql_t += " WHERE " + " AND ".join(where_t)
    sql_t += " ORDER BY t.id DESC"

    cur.execute(sql_t, params_t)
    trab = cur.fetchall()

    sql_titulos = """
        SELECT t.id, t.nome, t.numero_titulo, t.titulo_valido, t.titulo_uf, t.municipio,
               t.zona, t.secao, u.nome AS lideranca_nome
        FROM trabalho t
        JOIN usuarios u ON u.id=t.lideranca_id
    """
    params_titulos = []
    if titulo_status == "validos":
        sql_titulos += " WHERE t.titulo_valido=1"
    elif titulo_status == "invalidos":
        sql_titulos += " WHERE t.titulo_valido=0"
    elif titulo_status == "nao_informados":
        sql_titulos += " WHERE t.titulo_valido IS NULL"
    sql_titulos += " ORDER BY t.id DESC"
    cur.execute(sql_titulos, params_titulos)
    titulos_lista = cur.fetchall()

    cur.close()
    conn.close()

    return render_template(
        "admin.html",
        liderancas=liderancas,
        ranking=ranking,
        espontaneos=esp,
        trabalhos=trab,
        total_liderancas=total_liderancas,
        total_pendentes=total_pendentes,
        total_esp=total_esp,
        total_trab=total_trab,
        total_votos_previstos=total_votos_previstos,
        total_lideres_rua=total_lideres_rua,
        total_trabalho_bu=total_trabalho_bu,
        total_titulos_validos=total_titulos_validos,
        total_titulos_invalidos=total_titulos_invalidos,
        total_titulos_nao_informados=total_titulos_nao_informados,
        titulos_lista=titulos_lista,
        titulo_status=titulo_status,
        busca=busca,
        municipio=municipio,
        lideranca_id=lideranca_id
    )



@app.route("/admin/lideranca/<int:user_id>/estimativa", methods=["POST"])
@login_required
@admin_required
def salvar_estimativa_lideranca(user_id):
    try: votos=max(0,int(request.form.get('votos_estimados',0) or 0))
    except ValueError: votos=0
    conn=db(); cur=conn.cursor(); cur.execute("UPDATE usuarios SET votos_estimados=%s WHERE id=%s AND perfil='lideranca'",(votos,user_id)); conn.commit(); cur.close(); conn.close()
    flash("Estimativa da liderança atualizada.")
    return redirect(url_for("admin") + "#liderancas")

@app.route("/admin/aprovar/<int:user_id>")
@login_required
@admin_required
def aprovar(user_id):
    conn = db()
    cur = conn.cursor()
    cur.execute("UPDATE usuarios SET status='ativo' WHERE id=%s AND perfil='lideranca'", (user_id,))
    conn.commit()
    cur.close()
    conn.close()
    flash("Liderança aprovada.")
    return redirect(url_for("admin"))


@app.route("/admin/bloquear/<int:user_id>")
@login_required
@admin_required
def bloquear(user_id):
    conn = db()
    cur = conn.cursor()
    cur.execute("UPDATE usuarios SET status='bloqueado' WHERE id=%s AND perfil='lideranca'", (user_id,))
    conn.commit()
    cur.close()
    conn.close()
    flash("Liderança bloqueada.")
    return redirect(url_for("admin"))


@app.route("/admin/apagar_lideranca/<int:user_id>")
@login_required
@admin_required
def apagar_lideranca(user_id):
    conn = db()
    cur = conn.cursor()
    cur.execute("DELETE FROM usuarios WHERE id=%s AND perfil='lideranca'", (user_id,))
    conn.commit()
    cur.close()
    conn.close()
    flash("Liderança apagada com todos os cadastros vinculados.")
    return redirect(url_for("admin"))


@app.route("/admin/permissao_trabalho/<int:user_id>/<acao>")
@login_required
@admin_required
def permissao_trabalho(user_id, acao):
    valor = 1 if acao == "liberar" else 0
    conn = db()
    cur = conn.cursor()
    cur.execute("UPDATE usuarios SET pode_trabalho=%s WHERE id=%s AND perfil='lideranca'", (valor, user_id))
    conn.commit()
    cur.close()
    conn.close()
    flash("Permissão de trabalho atualizada.")
    return redirect(url_for("admin"))



@app.route("/admin/permissao_edicao/<tipo>/<int:item_id>/<acao>")
@login_required
@admin_required
def permissao_edicao(tipo, item_id, acao):
    if tipo not in {"espontaneo", "trabalho"} or acao not in {"liberar", "bloquear"}:
        abort(404)
    tabela = "espontaneos" if tipo == "espontaneo" else "trabalho"
    valor = 1 if acao == "liberar" else 0
    conn = db(); cur = conn.cursor()
    cur.execute(f"UPDATE {tabela} SET edit_liberado=%s WHERE id=%s", (valor, item_id))
    conn.commit(); cur.close(); conn.close()
    flash("Edição liberada para a liderança." if valor else "Edição bloqueada para a liderança.")
    return redirect(request.referrer or url_for("admin"))


@app.route("/lideranca/editar_espontaneo/<int:item_id>", methods=["GET", "POST"])
@login_required
@lideranca_required
def lideranca_editar_espontaneo(item_id):
    conn = db(); cur = conn.cursor()
    cur.execute("SELECT * FROM espontaneos WHERE id=%s AND lideranca_id=%s", (item_id, session["user_id"]))
    item = cur.fetchone()
    if not item or int(item.get("edit_liberado") or 0) != 1:
        cur.close(); conn.close(); flash("A edição deste contato ainda não foi liberada pelo administrador.")
        return redirect(url_for("lideranca"))
    if request.method == "POST":
        nome = request.form.get("nome_completo", "").strip()
        municipio = request.form.get("municipio", "").strip()
        telefone = request.form.get("telefone", "").strip()
        endereco = request.form.get("endereco_completo", "").strip()
        cur.execute("""UPDATE espontaneos SET nome_completo=%s, municipio=%s, telefone=%s, endereco_completo=%s,
                       nome_normalizado=%s, telefone_normalizado=%s, edit_liberado=0
                       WHERE id=%s AND lideranca_id=%s""",
                    (nome, municipio, telefone, endereco, normalizar(nome), normalizar(telefone), item_id, session["user_id"]))
        conn.commit(); cur.close(); conn.close(); flash("Contato atualizado. A edição foi fechada novamente.")
        return redirect(url_for("lideranca"))
    cur.close(); conn.close()
    return render_template("editar_espontaneo.html", item=item, modo_lideranca=True)


@app.route("/lideranca/editar_trabalho/<int:item_id>", methods=["GET", "POST"])
@login_required
@lideranca_required
def lideranca_editar_trabalho(item_id):
    conn = db(); cur = conn.cursor()
    cur.execute("SELECT pode_trabalho FROM usuarios WHERE id=%s", (session["user_id"],))
    user = cur.fetchone()
    if not user or int(user["pode_trabalho"] or 0) != 1:
        cur.close(); conn.close(); flash("Você não tem autorização para acessar o Cadastro 02.")
        return redirect(url_for("lideranca"))
    cur.execute("SELECT * FROM trabalho WHERE id=%s AND lideranca_id=%s", (item_id, session["user_id"]))
    item = cur.fetchone()
    if not item or int(item.get("edit_liberado") or 0) != 1:
        cur.close(); conn.close(); flash("A edição deste contato ainda não foi liberada pelo administrador.")
        return redirect(url_for("lideranca"))
    if request.method == "POST":
        nome = request.form.get("nome", "").strip(); municipio = request.form.get("municipio", "").strip()
        colegio = request.form.get("colegio", "").strip(); endereco = request.form.get("endereco", "").strip()
        telefone = request.form.get("telefone", "").strip(); zona = request.form.get("zona", "").strip()
        secao = request.form.get("secao", "").strip(); titulo = request.form.get("numero_titulo", "").strip()
        tipo_trabalho = request.form.get("tipo_trabalho", "Boca de Urna").strip()
        if tipo_trabalho not in ("Líder de rua", "Boca de Urna", "Trabalho de BU"): tipo_trabalho = "Boca de Urna"
        try: votos_previstos = max(0, int(request.form.get("votos_previstos", "0") or 0))
        except ValueError: votos_previstos = 0
        titulo_info = analisar_titulo_eleitor(titulo)
        foto_arquivo = request.files.get("foto")
        foto = foto_arquivo.read() if foto_arquivo and foto_arquivo.filename else None
        foto_mime = (foto_arquivo.mimetype or "image/jpeg") if foto_arquivo and foto_arquivo.filename else None
        if foto and len(foto) > 5 * 1024 * 1024:
            cur.close(); conn.close(); flash("A foto deve ter no máximo 5 MB.")
            return redirect(url_for("lideranca_editar_trabalho", item_id=item_id))
        cur.execute("""UPDATE trabalho SET nome=%s, municipio=%s, colegio=%s, endereco=%s, telefone=%s,
                       zona=%s, secao=%s, numero_titulo=%s, nome_normalizado=%s, telefone_normalizado=%s,
                       titulo_normalizado=%s, titulo_valido=%s, titulo_uf=%s, tipo_trabalho=%s,
                       votos_previstos=%s, foto=COALESCE(%s,foto), foto_mime=COALESCE(%s,foto_mime), edit_liberado=0
                       WHERE id=%s AND lideranca_id=%s""",
                    (nome, municipio, colegio, endereco, telefone, zona, secao, titulo, normalizar(nome),
                     normalizar(telefone), normalizar_titulo(titulo), None if titulo_info["valido"] is None else int(titulo_info["valido"]), titulo_info["uf"],
                     tipo_trabalho, votos_previstos, foto, foto_mime, item_id, session["user_id"]))
        conn.commit(); cur.close(); conn.close(); flash("Contato atualizado. A edição foi fechada novamente.")
        return redirect(url_for("lideranca"))
    cur.close(); conn.close()
    return render_template("editar_trabalho.html", item=item, modo_lideranca=True)


def gerar_docx(titulo, cabecalhos, linhas):
    documento = Document()
    secao = documento.sections[0]
    secao.orientation = WD_ORIENT.LANDSCAPE
    secao.page_width, secao.page_height = secao.page_height, secao.page_width
    secao.top_margin = Inches(0.45); secao.bottom_margin = Inches(0.45)
    secao.left_margin = Inches(0.45); secao.right_margin = Inches(0.45)
    p = documento.add_paragraph()
    r = p.add_run(titulo); r.bold = True; r.font.size = Pt(16)
    documento.add_paragraph(f"Gerado em {datetime.now().strftime('%d/%m/%Y às %H:%M')}")
    tabela = documento.add_table(rows=1, cols=len(cabecalhos))
    tabela.alignment = WD_TABLE_ALIGNMENT.CENTER
    tabela.style = "Table Grid"
    for i, texto in enumerate(cabecalhos):
        celula = tabela.rows[0].cells[i]; celula.text = str(texto); celula.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for run in celula.paragraphs[0].runs: run.bold = True; run.font.size = Pt(8)
    for linha in linhas:
        cells = tabela.add_row().cells
        for i, valor in enumerate(linha):
            cells[i].text = "" if valor is None else str(valor)
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            for run in cells[i].paragraphs[0].runs: run.font.size = Pt(8)
    buffer = io.BytesIO(); documento.save(buffer); buffer.seek(0); return buffer


@app.route("/admin/exportar_word/<tipo>")
@login_required
@admin_required
def exportar_word(tipo):
    lideranca_id = request.args.get("lideranca_id", "").strip()
    conn = db(); cur = conn.cursor()
    lideranca_nome = None

    if lideranca_id:
        cur.execute("SELECT id, nome FROM usuarios WHERE id=%s AND perfil='lideranca'", (lideranca_id,))
        lideranca = cur.fetchone()
        if not lideranca:
            cur.close(); conn.close()
            flash("Liderança não encontrada.")
            return redirect(url_for("admin") + "#relatorios")
        lideranca_nome = lideranca["nome"]

    if tipo == "espontaneos":
        sql = """SELECT u.nome AS lideranca, e.nome_completo, e.municipio, e.telefone, e.endereco_completo
                 FROM espontaneos e JOIN usuarios u ON u.id=e.lideranca_id"""
        params = []
        if lideranca_id:
            sql += " WHERE e.lideranca_id=%s"
            params.append(lideranca_id)
        sql += " ORDER BY e.id DESC"
        cur.execute(sql, params)
        rows = cur.fetchall(); headers = ["Liderança", "Nome", "Município", "Telefone", "Endereço"]
        linhas = [[r["lideranca"], r["nome_completo"], r["municipio"], r["telefone"], r["endereco_completo"]] for r in rows]
        titulo = "Contatos espontâneos"
        nome = "contatos_espontaneos_editavel.docx"
        if lideranca_nome:
            titulo += f" - Liderança: {lideranca_nome}"
            nome = f"espontaneos_lideranca_{lideranca_id}.docx"
    elif tipo == "trabalho":
        sql = """SELECT u.nome AS lideranca, t.nome, t.municipio, t.colegio, t.telefone, t.zona, t.secao,
                        t.numero_titulo, t.titulo_uf, t.tipo_trabalho, t.votos_previstos
                 FROM trabalho t JOIN usuarios u ON u.id=t.lideranca_id"""
        params = []
        if lideranca_id:
            sql += " WHERE t.lideranca_id=%s"
            params.append(lideranca_id)
        sql += " ORDER BY t.id DESC"
        cur.execute(sql, params)
        rows = cur.fetchall(); headers = ["Liderança", "Nome", "Município", "Colégio", "Telefone", "Zona", "Seção", "Título", "Origem", "Tipo", "Previsão"]
        linhas = [[r[h] for h in ["lideranca","nome","municipio","colegio","telefone","zona","secao","numero_titulo","titulo_uf","tipo_trabalho","votos_previstos"]] for r in rows]
        titulo = "Cadastros de trabalho"
        nome = "cadastros_trabalho_editavel.docx"
        if lideranca_nome:
            titulo += f" - Liderança: {lideranca_nome}"
            nome = f"trabalho_lideranca_{lideranca_id}.docx"
    elif tipo in {"titulos_validos", "titulos_invalidos"}:
        valor = 1 if tipo == "titulos_validos" else 0
        cur.execute("""SELECT u.nome AS lideranca, t.nome, t.numero_titulo, t.titulo_uf, t.municipio, t.zona, t.secao
                       FROM trabalho t JOIN usuarios u ON u.id=t.lideranca_id WHERE t.titulo_valido=%s ORDER BY t.id DESC""", (valor,))
        rows = cur.fetchall(); headers = ["Liderança", "Nome", "Título", "Origem", "Município", "Zona", "Seção"]
        linhas = [[r[h] for h in ["lideranca","nome","numero_titulo","titulo_uf","municipio","zona","secao"]] for r in rows]
        titulo = "Títulos válidos" if valor else "Títulos inválidos"; nome = ("titulos_validos" if valor else "titulos_invalidos") + ".docx"
    else:
        cur.close(); conn.close(); abort(404)
    cur.close(); conn.close()
    return send_file(gerar_docx(titulo, headers, linhas), as_attachment=True, download_name=nome,
                     mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")


def gerar_pdf(titulo, cabecalhos, linhas, larguras=None, orientacao=landscape(A4)):
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=orientacao, rightMargin=24, leftMargin=24, topMargin=28, bottomMargin=28)
    styles = getSampleStyleSheet()
    titulo_style = ParagraphStyle("TituloPainel", parent=styles["Title"], alignment=TA_CENTER, textColor=colors.HexColor("#071c4d"), fontSize=18, spaceAfter=14)
    elementos = [Paragraph(titulo, titulo_style), Spacer(1, 6)]
    dados = [cabecalhos] + [[str(v if v is not None else "") for v in linha] for linha in linhas]
    tabela = Table(dados, repeatRows=1, colWidths=larguras)
    tabela.setStyle(TableStyle([
        ("BACKGROUND", (0,0), (-1,0), colors.HexColor("#071c4d")),
        ("TEXTCOLOR", (0,0), (-1,0), colors.white),
        ("FONTNAME", (0,0), (-1,0), "Helvetica-Bold"),
        ("FONTSIZE", (0,0), (-1,-1), 7.5),
        ("GRID", (0,0), (-1,-1), .35, colors.HexColor("#d9e3f0")),
        ("ROWBACKGROUNDS", (0,1), (-1,-1), [colors.white, colors.HexColor("#f4f7fb")]),
        ("VALIGN", (0,0), (-1,-1), "TOP"),
        ("LEFTPADDING", (0,0), (-1,-1), 5),
        ("RIGHTPADDING", (0,0), (-1,-1), 5),
        ("TOPPADDING", (0,0), (-1,-1), 5),
        ("BOTTOMPADDING", (0,0), (-1,-1), 5),
    ]))
    elementos.append(tabela)
    doc.build(elementos)
    buffer.seek(0)
    return buffer


@app.route("/admin/exportar_pdf/<tipo>")
@login_required
@admin_required
def exportar_pdf(tipo):
    conn = db()
    cur = conn.cursor()
    if tipo == "espontaneos":
        cur.execute("""SELECT u.nome AS lideranca, e.nome_completo, e.municipio, e.telefone, e.endereco_completo
                       FROM espontaneos e JOIN usuarios u ON u.id=e.lideranca_id ORDER BY e.id DESC""")
        rows=cur.fetchall()
        linhas=[[r["lideranca"],r["nome_completo"],r["municipio"],r["telefone"],r["endereco_completo"]] for r in rows]
        pdf=gerar_pdf("Contatos espontâneos", ["Liderança","Nome","Município","Telefone","Endereço"], linhas, [105,135,90,90,270])
        nome="contatos_espontaneos.pdf"
    elif tipo == "trabalho":
        cur.execute("""SELECT u.nome AS lideranca, t.nome, t.municipio, t.colegio, t.telefone, t.zona, t.secao, t.numero_titulo, t.tipo_trabalho, t.votos_previstos
                       FROM trabalho t JOIN usuarios u ON u.id=t.lideranca_id ORDER BY t.id DESC""")
        rows=cur.fetchall()
        linhas=[[r["lideranca"],r["nome"],r["municipio"],r["colegio"],r["telefone"],r["zona"],r["secao"],r["numero_titulo"],r["tipo_trabalho"],r["votos_previstos"]] for r in rows]
        pdf=gerar_pdf("Contatos para trabalho", ["Liderança","Nome","Município","Colégio","Telefone","Zona","Seção","Título","Tipo","Previsão"], linhas, [70,82,62,82,66,34,34,66,72,45])
        nome="contatos_trabalho.pdf"
    else:
        cur.close(); conn.close(); flash("Relatório inválido."); return redirect(url_for("admin"))
    cur.close(); conn.close()
    return send_file(pdf, as_attachment=True, download_name=nome, mimetype="application/pdf")


@app.route("/admin/editar_espontaneo/<int:item_id>", methods=["GET", "POST"])
@login_required
@admin_required
def editar_espontaneo(item_id):
    conn = db()
    cur = conn.cursor()

    if request.method == "POST":
        nome = request.form.get("nome_completo", "").strip()
        municipio = request.form.get("municipio", "").strip()
        telefone = request.form.get("telefone", "").strip()
        endereco = request.form.get("endereco_completo", "").strip()

        cur.execute("""
            UPDATE espontaneos
            SET nome_completo=%s, municipio=%s, telefone=%s, endereco_completo=%s,
                nome_normalizado=%s, telefone_normalizado=%s
            WHERE id=%s
        """, (nome, municipio, telefone, endereco, normalizar(nome), normalizar(telefone), item_id))

        conn.commit()
        cur.close()
        conn.close()

        flash("Cadastro espontâneo editado.")
        return redirect(url_for("admin"))

    cur.execute("""
        SELECT e.*, u.nome AS lideranca_nome
        FROM espontaneos e
        JOIN usuarios u ON u.id=e.lideranca_id
        WHERE e.id=%s
    """, (item_id,))
    item = cur.fetchone()

    cur.close()
    conn.close()

    return render_template("editar_espontaneo.html", item=item)


@app.route("/admin/editar_trabalho/<int:item_id>", methods=["GET", "POST"])
@login_required
@admin_required
def editar_trabalho(item_id):
    conn = db()
    cur = conn.cursor()

    if request.method == "POST":
        nome = request.form.get("nome", "").strip()
        municipio = request.form.get("municipio", "").strip()
        colegio = request.form.get("colegio", "").strip()
        endereco = request.form.get("endereco", "").strip()
        telefone = request.form.get("telefone", "").strip()
        zona = request.form.get("zona", "").strip()
        secao = request.form.get("secao", "").strip()
        titulo = request.form.get("numero_titulo", "").strip()
        tipo_trabalho = request.form.get("tipo_trabalho", "Boca de Urna").strip()
        if tipo_trabalho not in ("Líder de rua", "Boca de Urna", "Trabalho de BU"):
            tipo_trabalho = "Boca de Urna"
        try:
            votos_previstos = max(0, int(request.form.get("votos_previstos", "0") or 0))
        except ValueError:
            votos_previstos = 0

        titulo_info = analisar_titulo_eleitor(titulo)
        foto_arquivo = request.files.get("foto")
        foto = foto_arquivo.read() if foto_arquivo and foto_arquivo.filename else None
        foto_mime = (foto_arquivo.mimetype or "image/jpeg") if foto_arquivo and foto_arquivo.filename else None
        if foto and len(foto) > 5 * 1024 * 1024:
            cur.close(); conn.close()
            flash("A foto deve ter no máximo 5 MB.")
            return redirect(url_for("editar_trabalho", item_id=item_id))
        cur.execute("""
            UPDATE trabalho
            SET nome=%s, municipio=%s, colegio=%s, endereco=%s, telefone=%s,
                zona=%s, secao=%s, numero_titulo=%s,
                nome_normalizado=%s, telefone_normalizado=%s, titulo_normalizado=%s,
                titulo_valido=%s, titulo_uf=%s, tipo_trabalho=%s, votos_previstos=%s,
                foto=COALESCE(%s, foto), foto_mime=COALESCE(%s, foto_mime)
            WHERE id=%s
        """, (nome, municipio, colegio, endereco, telefone, zona, secao, titulo, normalizar(nome), normalizar(telefone), normalizar_titulo(titulo),
              None if titulo_info["valido"] is None else int(titulo_info["valido"]), titulo_info["uf"], tipo_trabalho, votos_previstos, foto, foto_mime, item_id))

        conn.commit()
        cur.close()
        conn.close()

        flash("Cadastro de trabalho editado.")
        return redirect(url_for("admin"))

    cur.execute("""
        SELECT t.*, u.nome AS lideranca_nome
        FROM trabalho t
        JOIN usuarios u ON u.id=t.lideranca_id
        WHERE t.id=%s
    """, (item_id,))
    item = cur.fetchone()

    cur.close()
    conn.close()

    return render_template("editar_trabalho.html", item=item)


@app.route("/admin/exportar/<tipo>")
@login_required
@admin_required
def exportar(tipo):
    conn = db()
    cur = conn.cursor()
    output = io.StringIO()
    writer = csv.writer(output)

    if tipo == "espontaneos":
        writer.writerow(["Lideranca", "Nome completo", "Municipio", "Telefone", "Endereco completo", "Data"])
        cur.execute("""
            SELECT e.*, u.nome AS lideranca_nome
            FROM espontaneos e
            JOIN usuarios u ON u.id=e.lideranca_id
            ORDER BY e.id DESC
        """)
        rows = cur.fetchall()
        for r in rows:
            writer.writerow([r["lideranca_nome"], r["nome_completo"], r["municipio"], r["telefone"], r["endereco_completo"], r["created_at"]])
        filename = "espontaneos.csv"
    else:
        writer.writerow(["Lideranca", "Nome", "Municipio", "Colegio", "Endereco", "Telefone", "Zona", "Secao", "Numero titulo", "Tipo de trabalho", "Votos previstos", "Data"])
        cur.execute("""
            SELECT t.*, u.nome AS lideranca_nome
            FROM trabalho t
            JOIN usuarios u ON u.id=t.lideranca_id
            ORDER BY t.id DESC
        """)
        rows = cur.fetchall()
        for r in rows:
            writer.writerow([r["lideranca_nome"], r["nome"], r["municipio"], r["colegio"], r["endereco"], r["telefone"], r["zona"], r["secao"], r["numero_titulo"], r["tipo_trabalho"], r["votos_previstos"], r["created_at"]])
        filename = "trabalho.csv"

    cur.close()
    conn.close()

    return Response(
        output.getvalue(),
        mimetype="text/csv",
        headers={"Content-Disposition": f"attachment; filename={filename}"}
    )


_db_inicializado = False

def garantir_banco():
    global _db_inicializado
    if _db_inicializado:
        return
    init_db()
    _db_inicializado = True

@app.before_request
def preparar_banco_antes_da_requisicao():
    if request.endpoint in {"index", "static", "health"}:
        return
    garantir_banco()

@app.errorhandler(413)
def arquivo_muito_grande(_erro):
    flash("A imagem é muito grande. Envie uma foto de até 5 MB.")
    destino = "lideranca" if session.get("perfil") == "lideranca" else "admin" if session.get("perfil") == "admin" else "index"
    return redirect(url_for(destino))

@app.route("/health")
def health():
    return {"status": "ok"}, 200

if __name__ == "__main__":
    port = int(os.environ.get("PORT", 5000))
    app.run(host="0.0.0.0", port=port, debug=False)
